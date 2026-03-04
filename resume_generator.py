from docx import Document
from docx.shared import Pt

def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def _add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def generate_resume_docx(data, selected_experience, jd_skills, company, target_role, output_path):
    doc = Document()

    # Header
    name = data.get("name", "")
    location = data.get("location", "")
    email = data.get("email", "")
    phone = data.get("phone", "")
    linkedin = data.get("linkedin", "")

    title = doc.add_paragraph()
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(f"{location} | {email} | {phone} | {linkedin}")

    # Summary (simple ATS-friendly)
    _add_heading(doc, "PROFILE SUMMARY")
    summary = (
        f"Fresh graduate in Network Engineering with hands-on experience in IT operations and cybersecurity. "
        f"Experienced in incident triage, Microsoft 365 migration support, security monitoring (Wazuh SIEM), "
        f"and compliance evidence preparation (ISO/IEC 27001:2022)."
    )
    doc.add_paragraph(summary)

    # Skills (optional highlight based on JD)
    _add_heading(doc, "KEY SKILLS")
    if jd_skills:
        doc.add_paragraph("Matched keywords: " + ", ".join(jd_skills))
    else:
        doc.add_paragraph("Networking, cybersecurity monitoring, incident triage, technical documentation, Microsoft 365 support.")

    # Experience
    _add_heading(doc, "EXPERIENCE")
    for exp in selected_experience:
        doc.add_paragraph(f"{exp['role']} — {exp['company']} | {exp['location']} | {exp['start']} – {exp['end']}")
        for b in exp.get("bullets", []):
            _add_bullet(doc, b["text"])

    # Education
    _add_heading(doc, "EDUCATION")
    for edu in data.get("education", []):
        doc.add_paragraph(f"{edu['degree']} — {edu['school']} ({edu.get('graduation','')})")
        for h in edu.get("highlights", []):
            _add_bullet(doc, h["text"])

    # Certifications
    _add_heading(doc, "CERTIFICATIONS")
    for c in data.get("certifications", []):
        _add_bullet(doc, c)

    # Skills (full list at end for ATS)
    _add_heading(doc, "TECHNICAL SKILLS")
    skills = data.get("skills", {})
    for cat, items in skills.items():
        doc.add_paragraph(f"{cat.title()}: {', '.join(items)}")

    doc.save(output_path)
